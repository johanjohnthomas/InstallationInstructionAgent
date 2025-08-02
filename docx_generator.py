from docx import Document
import re

def _add_run_with_markdown(paragraph, text):
    """
    Adds a run to a paragraph, parsing basic markdown like **bold**.
    
    Args:
        paragraph: The docx paragraph object.
        text (str): The text to add, which may contain markdown.
    """
    # Split by bold markers
    parts = re.split(r'(\**.*?\**)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Add bolded text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Add regular text
            paragraph.add_run(part)

def generate_document_from_markdown(title: str, markdown_content: str, output_filename_base: str, output_format: str):
    """
    Generates a document from a markdown string.

    Args:
        title (str): The main title for the document.
        markdown_content (str): The raw markdown content from the AI.
        output_filename_base (str): The base name for the output file.
        output_format (str): The desired output format ('docx' or 'md').
    """
    if output_format == "md":
        # For Markdown, just prepend the title and save
        full_content = f"# {title}\n\n{markdown_content}"
        output_filename = f"{output_filename_base}.md"
        with open(output_filename, "w", encoding="utf-8") as f:
            f.write(full_content)
        print(f"Generated Markdown file: {output_filename}")

    elif output_format == "docx":
        # For DOCX, parse the markdown and build the document
        document = Document()
        document.add_heading(title, level=0)
        
        lines = markdown_content.split('\n')
        in_code_block = False

        for line in lines:
            stripped_line = line.strip()

            # Handle headers
            if stripped_line.startswith('#'):
                level = len(stripped_line) - len(stripped_line.lstrip('#'))
                heading_text = stripped_line.lstrip('# ').strip()
                document.add_heading(heading_text, level=min(level, 4))
            
            # Handle bullet points (unordered lists)
            elif stripped_line.startswith(('* ', '- ')):
                paragraph_text = stripped_line[2:]
                p = document.add_paragraph(style='List Bullet')
                _add_run_with_markdown(p, paragraph_text)

            # Handle numbered lists
            elif re.match(r'^\d+\.\s', stripped_line):
                paragraph_text = re.sub(r'^\d+\.\s', '', stripped_line)
                p = document.add_paragraph(style='List Number')
                _add_run_with_markdown(p, paragraph_text)

            # Handle plain paragraphs
            elif stripped_line:
                p = document.add_paragraph()
                _add_run_with_markdown(p, stripped_line)
        
        output_filename = f"{output_filename_base}.docx"
        document.save(output_filename)
        print(f"Generated DOCX file: {output_filename}")
        
    else:
        print(f"Unsupported output format: {output_format}. Please choose 'docx' or 'md'.")

# Keep the original function for backward compatibility if needed, but rename it.
# Or remove it if it's confirmed to be obsolete. For now, we'll keep it commented out.
# def generate_document_from_dict(...):
#     ...

if __name__ == "__main__":
    # Example usage for testing the new function
    sample_title = "Docker Installation Guide for Ubuntu 22.04"
    sample_markdown = """
## What is Docker?
Docker is a platform that enables developers to **develop, deploy, and run** applications in containers.

## Installation Guide
1. Update your `apt` package index:
   `sudo apt update`
2. Install Docker's dependencies:
   `sudo apt install ca-certificates curl gnupg`
3. Add Docker's official GPG key.

## Pros & Cons
- **Pros:**
  - Portability
  - Isolation
- **Cons:**
  - Learning curve

## Next Steps
* Verify Docker installation: `docker run hello-world`
* Explore Docker commands.
"""
    
    # Generate DOCX from Markdown
    generate_document_from_markdown(sample_title, sample_markdown, "test_guide_from_md", "docx")
    
    # Generate Markdown from Markdown
    generate_document_from_markdown(sample_title, sample_markdown, "test_guide_from_md", "md")




